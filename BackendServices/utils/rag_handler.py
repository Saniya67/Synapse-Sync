# import faiss
# import numpy as np
# from sentence_transformers import SentenceTransformer
# from openai import OpenAI
# import os
# import re
#
# class RAGPipeline:
#     def __init__(self, index_path: str = "vectordb/index.faiss"):
#         self.embedder = SentenceTransformer("all-MiniLM-L6-v2")  # small + fast
#         self.index_path = index_path
#         self.dimension = 384  # embedding size for this model
#         self.documents = []   # store chunks
#
#         # Create FAISS index (L2 similarity)
#         self.index = faiss.IndexFlatL2(self.dimension)
#
#         # LLM client (OpenAI/Azure – can swap later)
#         self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
#
#     def _chunk_text(self, text: str, chunk_size: int = 300):
#         # Split into smaller chunks
#         sentences = re.split(r'(?<=[.!?]) +', text)
#         chunks, current = [], ""
#         for sent in sentences:
#             if len(current) + len(sent) < chunk_size:
#                 current += " " + sent
#             else:
#                 chunks.append(current.strip())
#                 current = sent
#         if current:
#             chunks.append(current.strip())
#         return chunks
#
#     def add_document(self, text: str):
#         chunks = self._chunk_text(text)
#         embeddings = self.embedder.encode(chunks)
#         self.index.add(np.array(embeddings, dtype=np.float32))
#         self.documents.extend(chunks)
#
#     def query(self, question: str, top_k: int = 3) -> str:
#         q_emb = self.embedder.encode([question])
#         D, I = self.index.search(np.array(q_emb, dtype=np.float32), top_k)
#
#         # Retrieve top chunks
#         retrieved = [self.documents[i] for i in I[0] if i < len(self.documents)]
#
#         # Build context
#         context = "\n".join(retrieved)
#
#         # Send to LLM
#         try:
#             response = self.client.chat.completions.create(
#                 model="gpt-4o-mini",
#                 messages=[
#                     {"role": "system", "content": "You are a helpful assistant answering questions based on context."},
#                     {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {question}"}
#                 ],
#                 max_tokens=300
#             )
#             return response.choices[0].message.content.strip()
#         except Exception as e:
#             return f"AI error: {e}"
import json
import os
import re
import faiss
import numpy as np
import pandas as pd
from sentence_transformers import SentenceTransformer
from openai import OpenAI
from typing import List, Union
from docx import Document
from BackendServices.utils.text_extractor import extract_text


#Execution in second case reading static file
# class RAGPipeline:
#     def __init__(self, index_path: str = "vectordb/index.faiss"):
#         self.embedder = SentenceTransformer("all-MiniLM-L6-v2")  # fast embedding
#         self.dimension = 384
#         self.index_path = index_path
#         self.sentences: List[str] = []  # each row converted to a sentence
#
#         # FAISS index
#         self.index = faiss.IndexFlatL2(self.dimension)
#
#         # LLM client
#         self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
#
#     # -------------------------
#     # Helper: extract rows from DOCX tables or CSV/XLSX
#     # -------------------------
#     def parse_file(self, file_path: str) -> pd.DataFrame:
#         if file_path.endswith(".csv"):
#             return pd.read_csv(file_path)
#         elif file_path.endswith((".xls", ".xlsx")):
#             return pd.read_excel(file_path)
#         elif file_path.endswith(".docx"):
#             doc = Document(file_path)
#             for table in doc.tables:
#                 data = []
#                 keys = [cell.text.strip() for cell in table.rows[0].cells]
#                 for row in table.rows[1:]:
#                     values = [cell.text.strip() for cell in row.cells]
#                     data.append(dict(zip(keys, values)))
#                 if data:
#                     return pd.DataFrame(data)
#         else:
#             raise ValueError("Unsupported file type")
#         return pd.DataFrame()
#
#     # -------------------------
#     # Convert each row to a natural-language sentence
#     # -------------------------
#     def rows_to_sentences(self, df: pd.DataFrame) -> List[str]:
#         sentences = []
#         for _, row in df.iterrows():
#             sentence_parts = []
#             for col in df.columns:
#                 sentence_parts.append(f"{col}: {row[col]}")
#             sentence = ", ".join(sentence_parts)
#             sentences.append(sentence)
#         return sentences
#
#     # -------------------------
#     # Add document (table or text) to RAG
#     # -------------------------
#     def add_document(self, content: Union[str, pd.DataFrame]):
#         if isinstance(content, pd.DataFrame):
#             rows = self.rows_to_sentences(content)
#         else:
#             # treat as unstructured text
#             rows = self._chunk_text(content)
#
#         embeddings = self.embedder.encode(rows)
#         self.index.add(np.array(embeddings, dtype=np.float32))
#         self.sentences.extend(rows)
#
#     # -------------------------
#     # Simple chunker for unstructured text
#     # -------------------------
#     def _chunk_text(self, text: str, chunk_size: int = 300) -> List[str]:
#         sentences = re.split(r'(?<=[.!?]) +', text)
#         chunks, current = [], ""
#         for sent in sentences:
#             if len(current) + len(sent) < chunk_size:
#                 current += " " + sent
#             else:
#                 chunks.append(current.strip())
#                 current = sent
#         if current:
#             chunks.append(current.strip())
#         return chunks
#
#     # -------------------------
#     # Query RAG for a question
#     # -------------------------
#     def query(self, question: str, top_k: int = 10) -> dict:
#         if len(self.sentences) == 0:
#             return {"error": "No documents added"}
#
#         # Embed question
#         q_emb = self.embedder.encode([question])
#         D, I = self.index.search(np.array(q_emb, dtype=np.float32), top_k)
#
#         # Retrieve top sentences
#         retrieved = [self.sentences[i] for i in I[0] if i < len(self.sentences)]
#         context = "\n".join(retrieved)
#
#         # Build prompt for structured JSON table output
#         prompt = f"""
#             You are a helpful assistant. Using ONLY the following context, answer the question.
#             Return the results as a JSON array of rows with all columns preserved.
#
#             Context:
#             {context}
#
#             Question: {question}
#             Output format example:
#             {{
#               "query": "...",
#               "result_table": [
#                 {{"SerialNo": "L1234", "Model": "Dell XPS", "Owner": "Ravi", "Status": "Active", "PurchaseDate": "2023-01-15"}}
#               ]
#             }}
#             """
#
#         try:
#             response = self.client.chat.completions.create(
#                 model="gpt-4o-mini",
#                 messages=[
#                     {"role": "system", "content": "You are a helpful assistant answering tabular questions."},
#                     {"role": "user", "content": prompt}
#                 ],
#                 max_tokens=500
#             )
#             answer = response.choices[0].message.content.strip()
#             return answer
#         except Exception as e:
#             return {"query": question, "error": str(e)}


class RAGPipeline:
    def __init__(self, index_path: str = "vectordb/index.faiss"):
        self.embedder = SentenceTransformer("all-MiniLM-L6-v2")
        self.dimension = 384
        self.index_path = index_path
        self.sentences: List[str] = []

        # FAISS index
        self.index = faiss.IndexFlatL2(self.dimension)
        self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    # -------------------------
    # Parse table files into DataFrame
    # -------------------------
    def parse_file(self, file_path: str) -> pd.DataFrame:
        if file_path.endswith(".csv"):
            return pd.read_csv(file_path)
        elif file_path.endswith((".xls", ".xlsx")):
            return pd.read_excel(file_path)
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            for table in doc.tables:
                data = []
                keys = [cell.text.strip() for cell in table.rows[0].cells]
                for row in table.rows[1:]:
                    values = [cell.text.strip() for cell in row.cells]
                    data.append(dict(zip(keys, values)))
                if data:
                    return pd.DataFrame(data)
        return pd.DataFrame()

    # -------------------------
    # Convert each row to sentence
    # -------------------------
    def rows_to_sentences(self, df: pd.DataFrame) -> List[str]:
        sentences = []
        for _, row in df.iterrows():
            sentence_parts = [f"{col}: {row[col]}" for col in df.columns]
            sentences.append(", ".join(sentence_parts))
        return sentences

    # -------------------------
    # Split unstructured text
    # -------------------------
    def _chunk_text(self, text: str, chunk_size: int = 300) -> List[str]:
        sentences = re.split(r'(?<=[.!?]) +', text)
        chunks, current = [], ""
        for sent in sentences:
            if len(current) + len(sent) < chunk_size:
                current += " " + sent
            else:
                chunks.append(current.strip())
                current = sent
        if current:
            chunks.append(current.strip())
        return chunks

    # -------------------------
    # Add a file to RAG index
    # -------------------------
    def add_file(self, file_path: str):
        try:
            df = self.parse_file(file_path)
            if not df.empty:
                rows = self.rows_to_sentences(df)
            else:
                raise ValueError("Empty table")
        except:
            text = extract_text(file_path)
            rows = self._chunk_text(text)

        embeddings = self.embedder.encode(rows)
        self.index.add(np.array(embeddings, dtype=np.float32))
        self.sentences.extend(rows)

    # -------------------------
    # Query RAG
    # -------------------------
    def query(self, question: str, top_k: int = 20) -> dict:
        if len(self.sentences) == 0:
            return {"error": "No documents indexed"}

        q_emb = self.embedder.encode([question])
        D, I = self.index.search(np.array(q_emb, dtype=np.float32), top_k)
        retrieved = [self.sentences[i] for i in I[0] if i < len(self.sentences)]
        context = "\n".join(retrieved)

        prompt = f"""
You are a helpful assistant. Using ONLY the following context, answer the question.
Return results in structured JSON format.

Context:
{context}

Question: {question}
Output format example:
{{
  "query": "...",
  "result_table": [
    {{"SerialNo": "L1234", "Model": "Dell XPS", "Owner": "Ravi", "Status": "Active", "PurchaseDate": "2023-01-15"}}
  ]
}}
"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant answering tabular questions."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=800
            )

            return response.choices[0].message.content.strip()
        except Exception as e:
            return {"query": question, "error": str(e)}

